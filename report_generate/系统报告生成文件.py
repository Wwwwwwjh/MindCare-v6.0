import docx
from docx import document
from docx.shared import Pt
from docx.shared import RGBColor
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from mpz import mpz1
from mpl import mpl1
path = r'report模板.docx'  # word文档路径
doc = Document(path)  # 创建一个文档对象
tables = doc.tables

def change_font(run):
    font = run.font
    font.name = '华文仿宋'  # 字体名称
    font.size = Pt(14)  # 字体大小
    font.bold = True  # 加粗

def change_font2(run):
    font = run.font
    font.name = '华文仿宋'  # 字体名称
    font.size = Pt(14)  # 字体大小
    font.bold = True  # 加粗

def change_font3(run):
    font = run.font
    font.name = '华文仿宋'  # 字体名称
    font.size = Pt(12)  # 字体大小
    font.bold = True  # 加粗
#num_1 到 num_10分别为模型各项分值输出，name_in为姓名，sex_in为性别，year_in为年龄，heart_rate为心率，BOX为血氧含量，pressure为压力，textp为综合判断，textad为专家建议
def make_report(num_1,num_2,num_3,num_4,num_5,num_6,num_7,num_8,num_9,num_10,name_in,sex_in,year_in,heart_rate,BOX,pressure,textp,textad):
    mpz1(num_1,num_2,num_3,num_4,num_5)     #生成柱状图可能病症分析
    mpl1(num_6,num_7,num_8,num_9,num_10)  #生成极坐标图社会生活自理
    # 获取表格中所有内容对应的坐标
    tab = 0
    tables[tab].cell(1, 1).text = 'num_1'  #情绪认知
    tables[tab].cell(1, 2).text = 'num_2'  # 交往
    tables[tab].cell(1, 3).text = 'num_3'  # 感觉
    tables[tab].cell(1, 4).text = 'num_4'  # 社会生活自理
    tables[tab].cell(1, 5).text = 'num_5'  # 语言
    tables[tab].cell(3, 1).text = 'num_6' + '%' # 非典型孤独症
    tables[tab].cell(3, 2).text = 'num_7' + '%' # 自闭症
    tables[tab].cell(3, 3).text = 'num_8' + '%'  # 强迫症
    tables[tab].cell(3, 4).text = 'num_9' + '%' # 脆性X综合征
    tables[tab].cell(3, 5).text = 'num_10' + '%' # 语言发育障碍
    #修改表格填充内容的字体并居中
    for i, row in enumerate(tables[tab].rows):
        for j, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if int(run.text) >= 70:
                        #字体颜色设置为红色
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    if int(run.text) <= 40 :
                        #字体颜色设置为绿色
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    font = run.font
                    font.name = '华文仿宋'  # 字体名称
                    font.size = Pt(12)  # 字体大小
                    font.bold = True  # 加粗
                    #将字体居中
                    font.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 修改文字
    for par in doc.paragraphs:
        for run in par.runs:
            text = run.text
            if 'name' == text:
                run.text = run.text.replace('name', name_in)
                change_font(run)
            if 'sex' == text:
                run.text = run.text.replace('sex', sex_in)
                change_font(run)
            if 'year' == text:
                run.text = run.text.replace('year', year_in)
                change_font(run)
            if 'n1' == text:            #心率
                run.text = run.text.replace('n1', heart_rate)
                change_font2(run)
            if 'n2' == text:            #血氧含量
                run.text = run.text.replace('n2',BOX )
                change_font2(run)
            if 'n3' == text:            #压力
                run.text = run.text.replace('n3', pressure)
                change_font2(run)
            if 'text1' == text:            #综合判断
                run.text = run.text.replace('text1', textp)
                change_font3(run)
            if 'text2' == text:            #专家建议
                run.text = run.text.replace('text2', textad)
                change_font3(run)
            if 'image1' == text:
                run.text = run.text.replace('image1', '')
                run.add_picture(r'C:\Users\23154\PycharmProjects\1111\testz.png', width=Inches(7.15/2.54/1.5), height=Inches(5.63/2.54/1.15))
            if 'image' == text:
                run.text = run.text.replace('image', '')
                run.add_picture(r'C:\Users\23154\PycharmProjects\1111\testl.png', width=Inches(7.15/2.54/1.5), height=Inches(5.63/2.54/1.15))
            if 'image3' == text:
                run.text = run.text.replace('image3', '')
                run.add_picture(r'C:\Users\23154\PycharmProjects\1111\p3.png', width=Inches(7.15/2.54/1.5), height=Inches(5.63/2.54/1.15))
    path1 = r'report_res.docx'
    doc.save(path1)


