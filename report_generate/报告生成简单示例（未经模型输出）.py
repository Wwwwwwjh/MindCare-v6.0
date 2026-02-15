import docx
from docx import document
from docx.shared import Pt
from docx.shared import RGBColor
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from mpz import mpz1
from mpl import mpl1
path = r'report模板.docx'  # word文档路径
doc = Document(path)  # 创建一个文档对象
tables = doc.tables

mpz1(43,79,35,93,51)     #生成柱状图可能病症分析
mpl1(57,70,10,24,59)  #生成极坐标图社会生活自理
# 获取表格中所有内容对应的坐标
tab = 0
tables[tab].cell(1, 1).text = '43'          #情绪认知
tables[tab].cell(1, 2).text = '79'  # 交往
tables[tab].cell(1, 3).text = '35'  # 感觉
tables[tab].cell(1, 4).text = '93'  # 社会生活自理
tables[tab].cell(1, 5).text = '51'  # 语言
tables[tab].cell(3, 1).text = '57' + '%' # 非典型孤独症
tables[tab].cell(3, 2).text = '70' + '%' # 自闭症
tables[tab].cell(3, 3).text = '10' + '%'  # 强迫症
tables[tab].cell(3, 4).text = '24' + '%' # 脆性X综合征
tables[tab].cell(3, 5).text = '59' + '%' # 语言发育障碍
#修改表格填充内容的字体并居中
for i, row in enumerate(tables[tab].rows):
    for j, cell in enumerate(row.cells):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.text == '35' or run.text == '70%' or run.text == '43'or run.text == '51':
                    #字体颜色设置为红色
                    run.font.color.rgb = RGBColor(255, 0, 0)
                if run.text == '93'  or run.text == '10%':
                    #字体颜色设置为绿色
                    run.font.color.rgb = RGBColor(0, 128, 0)
                font = run.font
                font.name = '华文仿宋'  # 字体名称
                font.size = Pt(12)  # 字体大小
                font.bold = True  # 加粗
                #将字体居中
                font.alignment = WD_ALIGN_PARAGRAPH.CENTER
for par in doc.paragraphs:
    for run in par.runs:
        print(run.text)

def change_font(run):
    font = run.font
    font.name = '华文仿宋'  # 字体名称
    font.size = Pt(14)  # 字体大小
    font.bold = True  # 加粗

def change_font2(run):
    font = run.font
    font.name = '华文仿宋'  # 字体名称
    font.size = Pt(14)  # 字体大小
    font.bold = True  # 加粗

def change_font3(run):
    font = run.font
    font.name = '华文仿宋'  # 字体名称
    font.size = Pt(12)  # 字体大小
    font.bold = True  # 加粗
# 修改文字
for par in doc.paragraphs:
    for run in par.runs:
        text = run.text
        if 'name' == text:
            run.text = run.text.replace('name', '华大为')
            change_font(run)
        if 'sex' == text:
            run.text = run.text.replace('sex', '女')
            change_font(run)
        if 'year' == text:
            run.text = run.text.replace('year', '10')
            change_font(run)
        if 'n1' == text:            #心率
            run.text = run.text.replace('n1', '87')
            change_font2(run)
        if 'n2' == text:            #血氧含量
            run.text = run.text.replace('n2', '95')
            change_font2(run)
        if 'n3' == text:            #压力
            run.text = run.text.replace('n3', '57')
            change_font2(run)
        if 'text1' == text:            #综合判断
            textp = '，华小为有较大概率患有自闭症，其在反应速度、人际交往方面、情绪认知得分较低，可能存在缺陷但有正常的感知和自我认知能力。'
            run.text = run.text.replace('text1', textp)
            change_font3(run)
        if 'text2' == text:            #专家建议
            textad = '1.问诊期间侧重对问诊者反应速度、社会交往方面能力的考查，进一步判断相关能力强弱。2.应用行为分析（ABA）、模仿法（Modeling）等常见心理疗法，帮助问诊者学习某些薄弱缺乏的行为。'
            run.text = run.text.replace('text2', textad)
            change_font3(run)
        if 'image1' == text:
            run.text = run.text.replace('image1', '')
            run.add_picture(r'C:\Users\23154\PycharmProjects\1111\testz.png', width=Inches(7.15/2.54/1.5), height=Inches(5.63/2.54/1.15))
        if 'image' == text:
            run.text = run.text.replace('image', '')
            run.add_picture(r'C:\Users\23154\PycharmProjects\1111\testl.png', width=Inches(7.15/2.54/1.5), height=Inches(5.63/2.54/1.15))
        if 'image3' == text:
            run.text = run.text.replace('image3', '')
            run.add_picture(r'C:\Users\23154\PycharmProjects\1111\p3.png', width=Inches(7.15/2.54/1.5), height=Inches(5.63/2.54/1.15))




path1 = r'report_res.docx'
doc.save(path1)


